import os
import time
import copy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx2pdf import convert

from vault.bar import *
from vault.translator import *
import VaultConfigFactory as vcf


class VaultCore:
    @staticmethod
    def CreateDocument(vaultConfig: vcf.VaultConfigFactory, src: str, dest: str, langs: list[str]):
        try:
            print(f"Loading source document: {src}")
            doc = Document(src)
            util = GoogleTranslate()
            config = vaultConfig.vaultConfig
            waitingTime = float(config.waitingTimeMilis) / 1000

            # Create new document
            new_doc = Document()
            
            total_elements = len(doc.paragraphs) + len(doc.tables)
            processed_elements = 0

            if config.verbose:
                print(f"Starting translation from '{langs[0]}' to '{langs[1]}'...")
                print(f"Total elements to process: {total_elements}")

            if config.showProgress:
                showProgress(0, total_elements, prefix='Progress:', suffix='Complete', length=25)

            # Process paragraphs first
            for i, paragraph in enumerate(doc.paragraphs):
                if config.verbose and i % 10 == 0:
                    print(f"Processing paragraph {i+1}/{len(doc.paragraphs)}")
                
                VaultCore.process_paragraph(util, new_doc, paragraph, langs, config)
                processed_elements += 1
                time.sleep(waitingTime)
                
                if config.showProgress:
                    showProgress(processed_elements, total_elements, prefix='Progress:', suffix='Complete', length=25)

            # Process tables
            for i, table in enumerate(doc.tables):
                if config.verbose:
                    print(f"Processing table {i+1}/{len(doc.tables)}")
                
                VaultCore.process_table(util, new_doc, table, langs, config)
                processed_elements += 1
                time.sleep(waitingTime)
                
                if config.showProgress:
                    showProgress(processed_elements, total_elements, prefix='Progress:', suffix='Complete', length=25)

            # Save the document
            if config.verbose:
                print(f"Saving to: {dest}")
            
            if dest.lower().endswith('.pdf'):
                temp_docx = dest.rsplit('.', 1)[0] + '.docx'
                new_doc.save(temp_docx)
                convert(temp_docx, dest)
                os.remove(temp_docx)
            else:
                new_doc.save(dest)

            if config.verbose:
                print("Document successfully translated and saved!")

        except Exception as e:
            print(f"Error in CreateDocument: {e}")
            import traceback
            traceback.print_exc()
            raise

    @staticmethod
    def process_paragraph(util, new_doc, paragraph, langs, config):
        """Process and translate a paragraph"""
        try:
            # Skip empty paragraphs
            if not paragraph.text.strip() and not VaultCore.has_images(paragraph):
                new_doc.add_paragraph()
                return

            # Handle paragraphs with images
            if VaultCore.has_images(paragraph):
                VaultCore.handle_image_paragraph(util, new_doc, paragraph, langs, config)
            else:
                VaultCore.handle_text_paragraph(util, new_doc, paragraph, langs, config)
                
        except Exception as e:
            print(f"Error processing paragraph: {e}")
            # Add the original paragraph as fallback
            new_para = new_doc.add_paragraph(paragraph.text)

    @staticmethod
    def has_images(paragraph):
        """Check if paragraph contains images using a simpler method"""
        try:
            # Look for drawing elements in the XML
            xml_str = str(paragraph._element.xml)
            return 'pic:pic' in xml_str or 'w:drawing' in xml_str or 'wp:inline' in xml_str
        except:
            return False

    @staticmethod
    def handle_image_paragraph(util, new_doc, paragraph, langs, config):
        """Handle paragraph that contains images"""
        new_para = new_doc.add_paragraph()
        
        # Copy paragraph formatting
        VaultCore.copy_paragraph_formatting(paragraph, new_para)
        
        # For image paragraphs, we'll copy the entire paragraph XML
        # This preserves images but we lose text translation
        try:
            # Clone the entire paragraph element
            new_para._p = copy.deepcopy(paragraph._p)
            
            # Now try to translate any text in the paragraph while preserving images
            # This is complex, so for now we'll keep the original text for image paragraphs
            # and just log a message
            if config.verbose and paragraph.text.strip():
                print(f"Image paragraph found with text: {paragraph.text[:100]}... (text not translated to preserve images)")
                
        except Exception as e:
            print(f"Error handling image paragraph: {e}")
            # Fallback: just copy the text
            if paragraph.text.strip():
                translated_text = util.GetText(paragraph.text, src=langs[0], dest=langs[1])
                if translated_text and translated_text.strip():
                    new_para.add_run(translated_text)

    @staticmethod
    def handle_text_paragraph(util, new_doc, paragraph, langs, config):
        """Handle text-only paragraph with translation"""
        if not paragraph.text.strip():
            # Empty paragraph, just copy
            new_doc.add_paragraph()
            return

        # Translate the text
        translated_text = util.GetText(paragraph.text, src=langs[0], dest=langs[1])
        
        if translated_text and translated_text.strip() and translated_text != paragraph.text:
            new_para = new_doc.add_paragraph(translated_text)
            # Copy formatting
            VaultCore.copy_paragraph_formatting(paragraph, new_para)
            
            # Try to copy run formatting for simple paragraphs
            if len(paragraph.runs) == 1 and len(new_para.runs) == 1:
                VaultCore.copy_run_formatting(paragraph.runs[0], new_para.runs[0])
            elif len(paragraph.runs) > 1:
                # For multiple runs, we need more complex handling
                # For now, we'll use the basic approach
                pass
        else:
            # Fallback to original text
            new_para = new_doc.add_paragraph(paragraph.text)
            VaultCore.copy_paragraph_formatting(paragraph, new_para)

    @staticmethod
    def process_table(util, new_doc, table, langs, config):
        """Process and translate table content"""
        try:
            # Create new table with same dimensions
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            
            # Copy table style if possible
            try:
                new_table.style = table.style
            except:
                pass

            # Translate each cell
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    VaultCore.translate_cell_content(util, new_cell, cell, langs, config)
                    
        except Exception as e:
            print(f"Error processing table: {e}")

    @staticmethod
    def translate_cell_content(util, new_cell, source_cell, langs, config):
        """Translate content within a table cell"""
        try:
            # Clear default content
            for paragraph in new_cell.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            new_cell._element.clear_content()
            
            # Process paragraphs in cell
            for para in source_cell.paragraphs:
                if para.text.strip():
                    translated_text = util.GetText(para.text, src=langs[0], dest=langs[1])
                    if translated_text and translated_text.strip():
                        new_para = new_cell.add_paragraph(translated_text)
                        VaultCore.copy_paragraph_formatting(para, new_para)
                else:
                    new_cell.add_paragraph()
                    
        except Exception as e:
            print(f"Error translating cell content: {e}")
            # Fallback: copy original text
            new_cell.text = source_cell.text

    @staticmethod
    def copy_paragraph_formatting(source_para, target_para):
        """Copy paragraph formatting"""
        try:
            target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
            target_para.style = source_para.style
            
            # Copy spacing
            target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
            target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
            target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
            
            # Copy indentation
            target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
            target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
            target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
            
        except Exception as e:
            # Silently continue if formatting can't be copied
            pass

    @staticmethod
    def copy_run_formatting(source_run, target_run):
        """Copy run formatting"""
        try:
            target_run.bold = source_run.bold
            target_run.italic = source_run.italic
            target_run.underline = source_run.underline
            
            if source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
                
            if source_run.font.name:
                target_run.font.name = source_run.font.name
                
            if source_run.font.size:
                target_run.font.size = source_run.font.size
                
        except Exception as e:
            # Silently continue if formatting can't be copied
            pass