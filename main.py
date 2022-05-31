import argparse as args
import time

import docx

from vault.translator import *
from vault.bar import *


def CreateDocument(src, dest, langs):
    doc = docx.Document(src)
    styles = [p.style for p in doc.paragraphs]
    array = [p.text for p in doc.paragraphs]
    ok = True

    n = len(array)
    util = GoogleTranslate()

    doc = docx.Document()
    showProgress(0, n, prefix='Progress:', suffix='Complete', length=25)

    for i, line in enumerate(array):
        try:
            newLine = util.GetText(line, src=langs[0], dest=langs[1])
            doc.add_paragraph(newLine, style=styles[i])
            time.sleep(1)
            showProgress(i + 1, n, prefix='Progress:', suffix='Complete', length=25)
        except AttributeError:
            ok = False
            break

    if ok:
        doc.save(dest)
        print("\n Successfully translate document!")


if __name__ == "__main__":
    parser = args.ArgumentParser(description='Small tool that translate docx documents.')
    parser.add_argument("-i", "--input", help="The input file containing the original document.")

    parser.add_argument("-langs", "--langs", help="List of languages used for translation.", nargs="*")
    parser.add_argument("-o", "--output", help="The output file that contains the generated document.")

    args = parser.parse_args()
    src = args.input

    dest = args.output
    langs = args.langs
    CreateDocument(src, dest, langs)
